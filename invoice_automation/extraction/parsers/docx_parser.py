"""DOCX parser using python-docx."""

from invoice_automation.extraction.file_handler import FileInfo
from invoice_automation.extraction.parsers.base_parser import ParsedDocument, ParserStrategy
from invoice_automation.utils.exceptions import ParsingError


class DOCXParserStrategy(ParserStrategy):
	"""Extracts text from DOCX files using python-docx."""

	def supports(self, file_info: FileInfo) -> bool:
		return file_info.file_type == "DOCX"

	def parse(self, file_info: FileInfo) -> ParsedDocument:
		try:
			from docx import Document

			doc = Document(file_info.file_path)
			paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

			# Also extract from tables
			for table in doc.tables:
				for row in table.rows:
					cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
					if cells:
						paragraphs.append(" | ".join(cells))

			text = "\n".join(paragraphs)

			return ParsedDocument(
				text=text,
				page_count=1,
				parsing_method="python_docx",
			)

		except ImportError:
			raise ParsingError("python-docx not installed. Install with: pip install python-docx")
		except Exception as e:
			raise ParsingError(f"DOCX parsing failed: {e}", original=e) from e
